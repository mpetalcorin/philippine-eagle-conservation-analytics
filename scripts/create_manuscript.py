from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_ORIENTATION
from docx.shared import Twips
from pathlib import Path
import pandas as pd
import os, re, math

ROOT = Path(__file__).resolve().parents[1]
OUTDIR = ROOT / 'results'
FIGDIR = ROOT / 'figures' / 'manuscript'
DOCX = ROOT / 'manuscript' / 'Philippine_Eagle_Integrated_Conservation_Analytics_Manuscript.docx'

# ---------- helpers ----------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None:
            node = OxmlElement('w:'+m)
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_columns(section, num=2, space_twips=360):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num))
    if num > 1:
        cols.set(qn('w:space'), str(space_twips))
    else:
        if qn('w:space') in cols.attrib:
            del cols.attrib[qn('w:space')]

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def add_hyperlink(paragraph, text, url, color='0563C1', underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if underline:
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '15'); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)
    return hyperlink

def add_markup_paragraph(doc, text='', style=None, align=None, first_line=True, space_after=2, keep_with_next=False):
    p=doc.add_paragraph(style=style)
    if align is not None: p.alignment=align
    fmt=p.paragraph_format
    fmt.space_after=Pt(space_after)
    fmt.line_spacing=1.02
    if first_line: fmt.first_line_indent=Inches(0.18)
    fmt.keep_with_next=keep_with_next
    # parse **bold**, *italic*
    pos=0
    pattern=re.compile(r'(\*\*.*?\*\*|\*.*?\*)')
    for m in pattern.finditer(text):
        if m.start()>pos: p.add_run(text[pos:m.start()])
        token=m.group(0)
        if token.startswith('**'):
            r=p.add_run(token[2:-2]); r.bold=True
        else:
            r=p.add_run(token[1:-1]); r.italic=True
        pos=m.end()
    if pos<len(text): p.add_run(text[pos:])
    return p

def add_heading(doc, text, level=1):
    style='Heading 1' if level==1 else 'Heading 2'
    p=doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next=True
    p.paragraph_format.space_before=Pt(6 if level==1 else 4)
    p.paragraph_format.space_after=Pt(2)
    p.add_run(text)
    return p

def add_figure(doc, image_path, caption, width=Inches(6.95)):
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(); r.add_picture(str(image_path), width=width)
    cap=doc.add_paragraph(style='Caption')
    cap.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after=Pt(4)
    cap.paragraph_format.keep_with_next=False
    # bold Figure label, italic species names via markup not needed in captions? handle manually
    label, rest=caption.split('.',1)
    rr=cap.add_run(label+'.'); rr.bold=True
    rest=rest.strip()
    # simple italic parsing
    pos=0
    for m in re.finditer(r'\*.*?\*',rest):
        if m.start()>pos: cap.add_run(rest[pos:m.start()])
        ri=cap.add_run(m.group(0)[1:-1]); ri.italic=True
        pos=m.end()
    if pos<len(rest): cap.add_run(rest[pos:])
    return cap

def full_width_figure(doc, image_path, caption):
    sec=doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_columns(sec,1)
    add_figure(doc,image_path,caption,Inches(6.95))
    sec2=doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_columns(sec2,2,330)
    return sec2

def add_reference(doc, text, doi=None, note=None):
    p=doc.add_paragraph(style='References')
    p.paragraph_format.left_indent=Inches(0.15)
    p.paragraph_format.first_line_indent=Inches(-0.15)
    p.paragraph_format.space_after=Pt(2)
    p.add_run(text)
    if doi:
        p.add_run(' ')
        add_hyperlink(p, f'https://doi.org/{doi}', f'https://doi.org/{doi}')
    if note:
        p.add_run(' '+note)
    return p

def add_table_from_df(doc, df, title, col_widths=None, font_size=7.0, max_rows=None, formats=None):
    add_markup_paragraph(doc, title, style='Supplementary Caption', first_line=False, space_after=2, keep_with_next=True)
    if max_rows is not None:
        df=df.head(max_rows)
    table=doc.add_table(rows=1, cols=len(df.columns))
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.style='Table Grid'
    hdr=table.rows[0]
    set_repeat_table_header(hdr)
    for j,c in enumerate(df.columns):
        cell=hdr.cells[j]; cell.text=str(c)
        set_cell_shading(cell,'D9EAD3')
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.size=Pt(font_size)
    for _,row in df.iterrows():
        cells=table.add_row().cells
        for j,c in enumerate(df.columns):
            val=row[c]
            if pd.isna(val): txt=''
            elif formats and c in formats: txt=formats[c](val)
            elif isinstance(val,float): txt=f'{val:.3f}'
            else: txt=str(val)
            cells[j].text=txt
            cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[j])
            for p in cells[j].paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs: r.font.size=Pt(font_size)
    if col_widths:
        for row in table.rows:
            for j,w in enumerate(col_widths):
                row.cells[j].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return table

# ---------- document and styles ----------
doc=Document()
sec=doc.sections[0]
sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
sec.top_margin=Inches(0.55); sec.bottom_margin=Inches(0.55); sec.left_margin=Inches(0.62); sec.right_margin=Inches(0.62)
set_columns(sec,1)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Times New Roman'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); normal.font.size=Pt(8.5)
normal.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after=Pt(2)
normal.paragraph_format.line_spacing=1.02
for name,size,color in [('Heading 1',11,'17365D'),('Heading 2',9.5,'274E13')]:
    st=styles[name]; st.font.name='Arial'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.keep_with_next=True
# caption
cap=styles['Caption']; cap.font.name='Times New Roman'; cap._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); cap.font.size=Pt(7.7); cap.font.italic=False; cap.font.color.rgb=RGBColor(0,0,0)
cap.paragraph_format.first_line_indent=Inches(0)
# references style
if 'References' not in styles:
    st=styles.add_style('References',WD_STYLE_TYPE.PARAGRAPH)
else: st=styles['References']
st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); st.font.size=Pt(7.4)
st.paragraph_format.line_spacing=1.0
# supplementary caption
if 'Supplementary Caption' not in styles:
    st=styles.add_style('Supplementary Caption',WD_STYLE_TYPE.PARAGRAPH)
else: st=styles['Supplementary Caption']
st.font.name='Times New Roman'; st.font.size=Pt(8); st.font.bold=True; st.font.color.rgb=RGBColor.from_string('17365D')

# Headers and footers are configured after all sections have been created.

# ---------- title page / front matter ----------
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
r=p.add_run('An Integrated Literature-Benchmarked Synthetic Conservation Analytics Framework for the Critically Endangered Philippine Eagle '); r.bold=True; r.font.name='Arial'; r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string('17365D')
r2=p.add_run('('); r2.bold=True; r2.font.name='Arial'; r2.font.size=Pt(16); r2.font.color.rgb=RGBColor.from_string('17365D')
r3=p.add_run('Pithecophaga jefferyi'); r3.bold=True; r3.italic=True; r3.font.name='Arial'; r3.font.size=Pt(16); r3.font.color.rgb=RGBColor.from_string('17365D')
r4=p.add_run('): Linking Habitat Fragmentation, Reproductive Ecology, Population Viability, and Conservation Genomics'); r4.bold=True; r4.font.name='Arial'; r4.font.size=Pt(16); r4.font.color.rgb=RGBColor.from_string('17365D')

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('Mark Ihrwell R. Petalcorin, PhD'); r.bold=True; r.font.size=Pt(11)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(1)
r=p.add_run('aAidea Ltd, United Kingdom'); r.italic=True; r.font.size=Pt(9)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
r=p.add_run('Correspondence: Mark Ihrwell R. Petalcorin'); r.font.size=Pt(8)

# status banner
banner=doc.add_table(rows=1,cols=1); banner.alignment=WD_TABLE_ALIGNMENT.CENTER; banner.style='Table Grid'
cell=banner.cell(0,0); set_cell_shading(cell,'FFF2CC'); cell.text='METHODS AND SCENARIO STUDY. ALL TERRITORY, BREEDING, DEMOGRAPHIC, AND GENOMIC RECORDS ARE SYNTHETIC. NO REAL NEST LOCATIONS, TELEMETRY COORDINATES, INDIVIDUAL IDENTITIES, OR GENOTYPES ARE INCLUDED.'
for p in cell.paragraphs:
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.bold=True; r.font.name='Arial'; r.font.size=Pt(8)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

add_heading(doc,'Abstract',1)
abstract=("The Philippine eagle (*Pithecophaga jefferyi*) is a critically endangered, forest-dependent apex predator whose recovery is constrained by extensive space requirements, slow reproduction, avoidable adult mortality, continued habitat loss, and low genomic diversity. Conservation evidence is distributed across population assessment, telemetry, nesting biology, remote sensing, and molecular studies, which makes it difficult to evaluate interacting threats within one transparent analytical system. We developed a reproducible, literature-benchmarked synthetic conservation framework that links ecological state, reproduction, habitat change, stochastic demography, and conservation genomics without exposing sensitive field data. The simulation represented 392 breeding territories distributed as an illustrative scenario across Mindanao, Luzon, Samar, and Leyte, and was calibrated to 28,624 km² of area of habitat, 32% protected coverage, a median 95% home range of 68 km², a median 50% core range of 13 km², 79% of space-time outside the core, and 12.74 km mean nearest-neighbour spacing. We generated 3,528 biennial territory-year records from 2010 to 2026, including 2,950 active breeding attempts. Mean incubation duration was 58.0 days, the female incubation share was 74.1%, hatching success was 62.4%, and fledging success per attempt was 37.9%. A logistic model achieved a test ROC-AUC of 0.674, a Brier score of 0.225, and a fivefold cross-validated ROC-AUC of 0.656 ± 0.017. Scenario projections indicated 2035 habitat losses of 368 km² under integrated protection compared with 1,250 km² under a high-pressure future. In 40-year stochastic simulations, median adult females declined from 392 to 106 under current pressure, increased to 526 under integrated conservation, and fell to 30 under high pressure, with a 99.6% probability of crossing a 50-female warning threshold in the latter scenario. Synthetic genomic data reproduced a mean nuclear heterozygosity of 0.000309, a minimum of 0.000207, and 17 mitochondrial haplotypes among 32 individuals. Genomic pairing optimisation improved the mean synthetic diversity score by 23.9% relative to random plans. These results show how nest protection, landscape connectivity, adult survival, genomic management, and community-based action can be analysed as interacting components of one conservation system. The outputs are hypotheses and decision-support demonstrations, not empirical estimates or an official recovery plan.")
add_markup_paragraph(doc,abstract,first_line=False,space_after=4)

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
r=p.add_run('Keywords: '); r.bold=True
p.add_run('Philippine eagle; conservation genomics; population viability analysis; habitat fragmentation; reproductive ecology; machine learning; synthetic data; decision support')

# two-column body
sec2=doc.add_section(WD_SECTION_START.CONTINUOUS); set_columns(sec2,2,330)

# ---------- Introduction ----------
add_heading(doc,'Introduction',1)
intro_paras=[
"The Philippine eagle (*Pithecophaga jefferyi*) is an endemic forest raptor and a national symbol of the Philippines. Its biological architecture combines traits that make it ecologically dominant yet demographically fragile: large body size, extensive ranging requirements, a single-egg clutch, prolonged parental investment, late recruitment, and a long generation interval. The species persists on Mindanao, Luzon, Samar, and Leyte, but its distribution is increasingly partitioned by forest conversion, roads, settlement, and other human pressures. Historical nesting work on Mindanao reported a mean nearest-neighbour distance of approximately 12.74 km, illustrating the broad landscape needed to accommodate breeding pairs (Bueser et al., 2003). More recent range-wide modelling estimated approximately 392 breeding pairs, 28,624 km² of area of habitat, and only about 32% of that area within protected lands (Sutton et al., 2023).",
"Movement ecology strengthens the argument that conservation cannot be reduced to the nest tree or a small circular buffer. Adult Philippine eagles have a reported median 95% home range of 68 km² and a median 50% core range of 13 km². Approximately 79% of space-time use occurs outside the core, meaning that most daily ecological function, including hunting, transit, territorial surveillance, and exposure to anthropogenic hazards, occurs beyond the most intensively used centre of activity (Sutton et al., 2024). A landscape may therefore retain a nest yet lose the prey base, canopy continuity, or safe movement routes needed to sustain the adults that maintain that nest.",
"Reproductive biology further amplifies the consequences of adult mortality and habitat degradation. Detailed observation at Mount Sinaka recorded an incubation duration of approximately 58 days and a female contribution of about 74% of incubation time (Ibañez et al., 2003). The energetic and behavioural investment in a single egg and dependent juvenile means that reproductive replacement is slow. In long-lived raptors, population growth is typically more sensitive to adult survival than to modest changes in fecundity. Consequently, a gunshot, trap injury, electrocution, collision, or loss of a breeding adult can remove not only one animal but also years of future reproductive output and pair-specific territorial knowledge.",
"The habitat problem is similarly multiscalar. Forest loss in Philippine Key Biodiversity Areas remains detectable over contemporary monitoring periods, and corridor-level analyses show that apparent stability at broad spatial scales can conceal local conversion and fragmentation (Daipan, 2021; Perez et al., 2020). Fragmentation changes more than the number of hectares. It increases edge exposure, alters canopy microclimate, changes prey communities, isolates breeding territories, and can force predators to travel farther through riskier matrices. Across tropical systems, primary and structurally intact forests support biological functions that degraded forests cannot fully replace, while persistent fragmentation produces long-lived effects on biodiversity and ecosystem dynamics (Gibson et al., 2011; Haddad et al., 2015; Watson et al., 2018).",
"Conservation also operates at the molecular scale. Population contraction and geographic isolation reduce effective population size, increase genetic drift, elevate the probability of mating among relatives, and permit long runs of homozygosity to accumulate. Homozygosity can expose recessive deleterious variants, reduce genome-wide allelic diversity, and constrain adaptive responses to pathogens and environmental change. These molecular processes can influence fertility, embryonic survival, immune competence, stress tolerance, and long-term evolutionary resilience. Recent Philippine eagle studies reported mitochondrial nucleotide diversity of approximately 0.00054 among 32 sampled individuals, with 17 haplotypes, and mean genome-wide nuclear heterozygosity of approximately 0.000309 among 35 individuals, with a minimum near 0.000207 (Bacus et al., 2025; Perdon et al., 2026). Mitochondrial haplotypes describe maternal lineage diversity, whereas nuclear heterozygosity and relatedness capture biparental genomic variation. These measures are complementary and should not be conflated.",
"Modern conservation science increasingly integrates remote sensing, field monitoring, demographic models, and genomics. Machine learning can help process expanding ecological datasets, but its value depends on ecological grounding, careful validation, explicit uncertainty, and protection against spatial or temporal leakage (Roberts et al., 2017; Tuia et al., 2022). Population viability analysis can compare management scenarios and identify sensitive life-history parameters, although it should be interpreted as conditional forecasting rather than deterministic prediction, particularly when empirical demographic series are short (Brook et al., 2000). Conservation genomics can improve pedigree reconstruction, kinship assessment, founder representation, and genetic rescue planning, but genomic findings must be translated through behavioural, veterinary, welfare, and ecological constraints (Shafer et al., 2015; Supple & Shapiro, 2018).",
"The present study addresses the need for a transparent integrative framework. We created a fully synthetic dataset calibrated to published Philippine eagle benchmarks and used it to test an end-to-end conservation workflow. The objectives were to reproduce key population, habitat, movement, breeding, forest-loss, and genomic central values; quantify relationships encoded between habitat condition and simulated breeding outcomes; compare habitat and population trajectories under alternative conservation scenarios; demonstrate lineage-aware pairing logic; identify ecological archetypes; and prioritise interventions under an illustrative budget. Because every individual-level record is simulated, the framework can be shared, audited, and modified without disclosing sensitive nest locations or real genotypes."
]
for x in intro_paras: add_markup_paragraph(doc,x)

# Figure 1 wide
full_width_figure(doc,Path('/mnt/data/philippine_eagle_guardian_of_rainforests.png'),
"Figure 1. Biological and conservation overview of the Philippine eagle (*Pithecophaga jefferyi*). The infographic summarises morphology, range, space requirements, reproductive biology, apex-predator function, conservation status, and public stewardship. The evolutionary-age panel present in an earlier version was intentionally removed. The figure is an educational synthesis and should not replace primary ecological measurements.")

# ---------- Methods ----------
add_heading(doc,'Materials and Methods',1)
add_heading(doc,'Study design, evidence selection, and safeguards',2)
methods_paras=[
"The analysis was designed as a literature-benchmarked synthetic study. Numerical targets were extracted from peer-reviewed journal articles indexed in PubMed, FAO AGRIS, or journal resources represented in USDA AGRICOLA. The core Philippine eagle registry included estimates of breeding-pair abundance, area of habitat, protected-area coverage, home range, core range, space-time outside core, nearest-neighbour spacing, incubation duration, female incubation share, forest loss, mitochondrial nucleotide diversity, haplotype richness, and genome-wide heterozygosity. The benchmark registry, DOI, uncertainty notation, and index provenance were retained in the notebook and exported as a machine-readable file.",
"No real nest coordinates, telemetry tracks, locality polygons, individual identities, pedigrees, or nucleotide sequences were used. Territory identifiers such as MI-001 were artificial codes. Island allocations were scenario values constrained to sum to the published global pair estimate and area of habitat, and must not be interpreted as island-specific censuses. The workflow is suitable for methods development, teaching, and sensitivity analysis, but not for law enforcement, nest visitation, or public geolocation of eagles.",
"The computational design followed reproducibility principles previously used by the author in synthetic molecular, genomic, and clinical modelling, in which benchmark registries, fixed random seeds, transparent assumptions, validation summaries, and exported intermediate datasets allow every result to be audited (Petalcorin, 2026a; Petalcorin, 2026b; Petalcorin, 2026c). Those publications are preprints and are cited only as methodological continuity, not as peer-reviewed evidence for Philippine eagle biology."
]
for x in methods_paras: add_markup_paragraph(doc,x)

add_heading(doc,'Synthetic territory generation and ecological variables',2)
methods2=[
"The model generated 392 breeding territories using a fixed random seed of 20260716. The illustrative allocation comprised 231 territories on Mindanao, 80 on Luzon, 47 on Samar, and 34 on Leyte. Corresponding areas of habitat were 16,888, 5,725, 3,435, and 2,576 km², respectively, summing to 28,624 km². For each territory, forest integrity was sampled from an island-adjusted beta distribution. Fragmentation was coupled inversely to forest integrity with stochastic variation. Human pressure increased with fragmentation, while canopy integrity and prey availability increased with forest condition and decreased with pressure. Distance to roads followed a gamma distribution modified by pressure, and persecution risk combined human pressure, fragmentation, and stochastic noise.",
"Exactly 32% of territories were assigned to protected status, with selection weighted toward higher forest integrity and lower human pressure. The 95% home range was generated from a log-normal function of fragmentation and prey limitation, then rescaled to a median of 68 km² and clipped to 39–161 km². Core range was modelled as a fraction of home range and rescaled to a median of 13 km², with bounds of 9–33 km². Space-time outside core was calibrated to a median of 79%, and nearest-neighbour spacing was calibrated to a mean of 12.74 km. A composite risk score integrated fragmentation, human pressure, persecution, prey limitation, canopy degradation, and lack of protection. Risk classes were defined as low, moderate, high, or critical using fixed score intervals.",
"This construction encoded a mechanistic hypothesis: prey depletion and canopy discontinuity increase ranging demand, wider movement increases contact with roads and human-dominated land, and higher exposure contributes to persecution and mortality risk. Because these dependencies were specified by design, associations recovered from the synthetic data demonstrate workflow behaviour rather than biological effect sizes."
]
for x in methods2: add_markup_paragraph(doc,x)

add_heading(doc,'Breeding histories and fledging-risk modelling',2)
methods3=[
"Every territory was evaluated at two-year intervals from 2010 through 2026, giving nine territory-year observations and 3,528 total records. A latent pair-quality term was assigned to each territory. Breeding-attempt probability increased with forest integrity, prey availability, and pair quality, and decreased with human pressure. Every active attempt contained one egg. Incubation duration was sampled around 58 days, and the female incubation share was drawn from a beta distribution calibrated to 74%.",
"Conditional hatching probability increased with forest integrity and prey availability, and decreased with human pressure, persecution risk, and the absolute magnitude of a synthetic climate anomaly. Conditional fledging probability increased with canopy integrity, prey availability, and pair quality, and decreased with pressure, persecution, and climate anomaly. The primary machine-learning endpoint was fledging among active attempts. Predictors were forest integrity, fragmentation, human pressure, canopy integrity, prey index, persecution risk, protected status, climate anomaly, and island.",
"Numeric predictors were standardised, island was one-hot encoded, and a class-weighted logistic regression model was fitted. Data were divided into 75% training and 25% test sets using stratification and the fixed seed. Discrimination was measured with the receiver operating characteristic area under the curve, calibration with the Brier score and calibration curve, and generalisability with fivefold stratified cross-validation. Standardised coefficients were reported for interpretability. The model was deliberately modest and transparent; it was not tuned to maximise performance because the objective was to demonstrate an auditable ecological workflow."
]
for x in methods3: add_markup_paragraph(doc,x)

add_heading(doc,'Forest-loss scenarios',2)
methods4=[
"The baseline annual loss rate was derived from 14,213 ha per year relative to an illustrative 4.5-million-ha forest baseline, approximately 0.316% annually. Territory-specific rates were modified by human pressure and fragmentation, with log-normal noise and bounds of 0.02–1.8% annually. Five multiplicative scenarios were evaluated from 2026 to 2035: current pressure, enforcement, corridor restoration, integrated protection, and a high-pressure future. The multipliers were 1.00, 0.72, 0.60, 0.42, and 1.45, respectively. Remaining area of habitat was calculated by compound annual retention at each territory and aggregated across the synthetic range."
]
for x in methods4: add_markup_paragraph(doc,x)

add_heading(doc,'Stochastic population viability analysis',2)
methods5=[
"A female-based stage model began with 392 adult females and four pre-breeding age classes. For each scenario, 250 stochastic trajectories were simulated for 40 years. Annual processes included binomial adult survival, breeding participation, nest success, an assumed 1:1 sex ratio among fledglings, and pre-breeder survival through four age classes. The scenarios were current pressure, habitat protection, anti-persecution, integrated conservation, and high pressure. Adult survival ranged from 0.915 to 0.958, breeding fraction from 0.46 to 0.59, nest success from 0.46 to 0.58, and pre-breeder survival from 0.66 to 0.76.",
"The model was not fitted to a long-term Philippine eagle mark–recapture series. Therefore, the absolute trajectories should be treated as scenario demonstrations. The principal comparison was the direction and relative magnitude of change when adult survival, breeding, nest success, and recruitment were altered together. Outcomes included the median adult-female abundance at year 40, 95% simulation intervals, the probability of falling below 50 adult females at any point, and the probability of extinction."
]
for x in methods5: add_markup_paragraph(doc,x)

add_heading(doc,'Synthetic conservation genomics and pairing optimisation',2)
methods6=[
"The nuclear dataset comprised 35 synthetic individuals with a mean heterozygosity constrained to 0.000309 and a minimum of 0.000207. A runs-of-homozygosity proxy decreased as heterozygosity increased, while two synthetic kinship axes represented multivariate relatedness structure without simulating real variants. The mitochondrial dataset comprised 32 synthetic individuals assigned to 17 haplotypes, with a mean nucleotide-diversity proxy of 0.00054. Haplotype diversity was calculated from haplotype frequencies.",
"A teaching example of pairing optimisation separated 18 synthetic females and 17 males. Pair scores combined distance in the two kinship dimensions, a small reward for cross-island representation, and the sum of individual heterozygosities. The Hungarian assignment algorithm maximised the total score subject to one-to-one pairing. The optimised mean was compared with 300 random pairing plans. This score is not a breeding recommendation. Real management would require genomic relatedness, pedigree validation, demographic representation, veterinary fitness, behavioural compatibility, reproductive history, welfare review, and explicit assessment of risks associated with moving individuals or mixing differentiated populations.",
"At the molecular level, genome-wide heterozygosity provides a broad measure of allelic variation across nuclear chromosomes, whereas runs of homozygosity indicate chromosomal segments inherited identically from related ancestors. Long homozygous tracts can reveal recent inbreeding and increase the probability that recessive alleles are expressed in the homozygous state. Mitochondrial haplotypes are inherited maternally and capture a different component of population history. Preserving rare maternal haplotypes cannot compensate for high nuclear relatedness, and maximising nuclear diversity cannot justify loss of unique maternal lineages."
]
for x in methods6: add_markup_paragraph(doc,x)

add_heading(doc,'Ecological archetypes and intervention portfolio',2)
methods7=[
"Six territory variables, forest integrity, fragmentation, human pressure, prey index, persecution risk, and annual forest-loss rate, were standardised and clustered with K-means using four clusters and 20 initialisations. Principal component analysis provided a two-dimensional visualisation. Clusters were labelled after examining their mean ecological profiles as relatively secure forest, watch-list landscape, fragmented pressure front, or critical intervention zone.",
"Each territory was assigned a recommended action by transparent rules. High persecution risk or proximity to roads triggered community patrol and anti-persecution action; high fragmentation triggered corridor restoration; low prey index triggered prey and habitat recovery; unprotected status triggered a stewardship or protected-area agreement; remaining sites received long-term monitoring. Each action had an illustrative cost and expected proportional risk reduction. Territories were ranked by benefit per cost, with an additional weight for sites outside protected areas, and selected until a budget of 300 arbitrary cost units was reached. The exercise demonstrates auditable prioritisation, not actual budgeting."
]
for x in methods7: add_markup_paragraph(doc,x)

add_heading(doc,'Software, execution, and statistical interpretation',2)
methods8=[
"The notebook used Python, NumPy, pandas, SciPy, Matplotlib, and scikit-learn. Every synthetic dataset, benchmark table, scenario output, and figure was exported. Numerical calibration was assessed by comparing synthetic and published central values and calculating relative error. Spearman correlation quantified the monotonic relationship between fragmentation and simulated home-range demand. Statistical P values generated from synthetic data describe internal consistency of the simulation and must not be interpreted as evidence from wild eagles.",
"The notebook was executed end to end before manuscript preparation. The complete analysis contains 392 synthetic territories, 3,528 breeding records, five forest scenarios, 1,250 population trajectories, 35 nuclear-genomic profiles, 32 mitochondrial profiles, 17 optimised pairings, four ecological archetypes, and an illustrative intervention portfolio."
]
for x in methods8: add_markup_paragraph(doc,x)

# Figure 2
full_width_figure(doc,FIGDIR/'figure_2_ecological_benchmarks.png',
"Figure 2. Literature-benchmarked synthetic population and landscape structure. (A) Illustrative allocation of 392 breeding territories across four islands. (B) Synthetic 95% home-range distributions calibrated to a median of 68 km². (C) Positive association between fragmentation and simulated ranging demand. (D) Protection gap by island, with the total protected proportion constrained to approximately 32%. Island allocations are scenarios and are not island-specific census estimates.")

# ---------- Results ----------
add_heading(doc,'Results',1)
add_heading(doc,'Benchmark fidelity and ecological state',2)
results1=[
"The synthetic dataset exactly reproduced the target totals of 392 breeding pairs and 28,624 km² of area of habitat. Protected coverage was 31.888%, a relative error of 0.351% from the 32% target. Median 95% home range, median 50% core range, median space-time outside core, and mean nearest-neighbour spacing reproduced their target values of 68 km², 13 km², 79%, and 12.74 km, respectively. These checks confirmed numerical calibration without claiming that the underlying island allocation or multivariable distributions represented observed population structure.",
"Of the 392 territories, 166 were classified as low risk, 136 as moderate risk, 75 as high risk, and 15 as critical risk. The simulated relationship between fragmentation and home-range demand was positive, as encoded by the generative model. Territories with low forest integrity and prey availability, high pressure, and greater persecution risk tended to require larger ranges and received higher composite risk scores. The pattern illustrates a potential exposure feedback: habitat degradation can enlarge the area through which adults must move, thereby increasing the number of roads, settlements, farms, and conflict interfaces encountered during foraging.",
"The protection-gap analysis showed that legal designation and ecological security were not equivalent. Protected territories were preferentially sampled from higher-integrity landscapes, but substantial areas of synthetic habitat and movement space remained outside protection. Because the median eagle spent 79% of its space-time outside the 50% core range, a nest-centred reserve would cover only a minority of daily ecological use."
]
for x in results1: add_markup_paragraph(doc,x)

add_heading(doc,'Breeding outcomes and predictive model',2)
results2=[
"The nine biennial assessments generated 2,950 active breeding attempts among 3,528 territory-year records. Each active attempt contained one egg. Mean incubation duration was 58.006 days and the mean female incubation share was 74.061%, closely reproducing the breeding benchmarks. Hatching occurred in 62.373% of active attempts, and 37.932% of attempts produced a fledgling. Because the hatching and fledging functions were synthetic, these percentages are scenario outputs rather than empirical estimates of wild productivity.",
"The logistic model produced a test ROC-AUC of 0.674 and a Brier score of 0.225. Fivefold cross-validated ROC-AUC was 0.656 ± 0.017. The model therefore separated higher- and lower-probability fledging outcomes better than chance, but retained substantial uncertainty, which is appropriate for a multifactorial ecological outcome. Calibration was broadly informative but imperfect.",
"The largest standardised negative coefficient was human pressure, -0.319, followed by fragmentation, -0.088. Persecution risk, -0.053, and climate anomaly, -0.036, were also negative. The largest positive coefficient was prey index, 0.161, followed by forest integrity, 0.064, canopy integrity, 0.054, and protected status, 0.040. Island coefficients were small after adjustment for the ecological predictors. These directions recovered the encoded rules, demonstrating that an interpretable model could trace how habitat, prey, and anthropogenic stress combine in a synthetic reproductive-risk system."
]
for x in results2: add_markup_paragraph(doc,x)

full_width_figure(doc,FIGDIR/'figure_3_breeding_model.png',
"Figure 3. Synthetic breeding outcomes and interpretable fledging-risk model. (A) Biennial counts of attempts, hatching, and fledging from 2010 to 2026. (B) Test-set receiver operating characteristic curve, ROC-AUC 0.674. (C) Standardised logistic-regression coefficients, with prey and habitat variables generally positive and human pressure variables negative. (D) Calibration curve, with Brier score 0.225. Associations validate the simulation workflow and are not empirical causal estimates.")

add_heading(doc,'Forest-loss projections',2)
results3=[
"Projected loss from the 2026 synthetic area of habitat diverged progressively among scenarios. By 2035, current pressure produced a cumulative loss of 868.0 km², leaving 27,756.0 km². Enforcement reduced the projected loss to 627.7 km², corridor restoration to 524.1 km², and integrated protection to 367.9 km². The high-pressure future produced the greatest loss, 1,249.7 km², leaving 27,374.3 km².",
"The difference between integrated protection and high pressure was 881.8 km² by 2035. This contrast arose from small annual differences compounded across many territories. The model therefore demonstrates why sustained reductions in local loss rates can produce substantial landscape benefits even when annual percentages appear modest."
]
for x in results3: add_markup_paragraph(doc,x)

add_heading(doc,'Population viability scenarios',2)
results4=[
"Starting from 392 adult females, the current-pressure scenario declined to a median of 106 adult females after 40 years, with a 95% simulation interval of 83.0–132.8. Habitat protection produced a median of 191, interval 155.2–239.8, and anti-persecution produced a median of 249, interval 205.2–302.3. Integrated conservation increased the median to 526, interval 437.2–603.6. The high-pressure future declined to a median of 30, interval 17.5–45.6.",
"No scenario reached simulated extinction within the 40-year horizon, but 99.6% of high-pressure trajectories crossed the warning threshold of 50 adult females. The other four scenarios did not cross that threshold in the 250 replicates. The strong benefit of anti-persecution relative to habitat-only protection reflected the high elasticity of a slow-breeding, long-lived species to adult survival. The integrated scenario performed best because it simultaneously improved adult survival, breeding participation, nest success, and pre-breeder survival."
]
for x in results4: add_markup_paragraph(doc,x)

full_width_figure(doc,FIGDIR/'figure_4_scenarios_pva.png',
"Figure 4. Habitat and demographic scenario projections. (A) Area-of-habitat trajectories from 2026 to 2035 under five pressure multipliers. (B) Median adult-female trajectories from 250 stochastic simulations per scenario over 40 years. (C) Probability of crossing the 50-adult-female warning threshold. The population model is comparative and assumption-driven, not a fitted forecast of the wild population.")

add_heading(doc,'Conservation genomics and pairing demonstration',2)
results5=[
"The synthetic nuclear-genomic dataset reproduced a mean heterozygosity of 0.000309 and a minimum of 0.000207 among 35 individuals. The mitochondrial dataset contained 17 haplotypes among 32 individuals and reproduced a mean nucleotide-diversity proxy of 0.000540. Calculated haplotype diversity was 0.9476, reflecting the retention of multiple maternal lineages despite low absolute nucleotide diversity.",
"The optimised assignment paired 17 female–male combinations. The mean synthetic pair score was 1.983, compared with a mean of 1.600 across random plans, a 23.9% improvement. Several high-scoring pairs combined individuals assigned to different island labels, but geographic mixing was only a small component of the score. The result shows that explicit optimisation can outperform ad hoc pairing under a defined objective, while also showing why the objective function must be scientifically and ethically justified before real use."
]
for x in results5: add_markup_paragraph(doc,x)

full_width_figure(doc,FIGDIR/'figure_5_genomics.png',
"Figure 5. Synthetic conservation-genomic outputs. (A) Genome-wide heterozygosity distribution calibrated to the published mean and minimum. (B) Frequencies of 17 mitochondrial haplotypes among 32 synthetic individuals. (C) Distribution of random pairing-plan scores compared with the optimised plan. No real sequences, pedigrees, relatedness values, or breeding recommendations are represented.")

add_heading(doc,'Conservation archetypes and intervention portfolio',2)
results6=[
"Clustering separated four ecological archetypes. The relatively secure forest cluster contained 119 territories, mean forest integrity 0.773, fragmentation 0.201, pressure 0.279, prey index 0.716, persecution risk 0.234, 57.1% protected coverage, and median risk 19.8. The watch-list cluster contained 105 territories and median risk 38.5. The fragmented pressure front contained 103 territories, mean fragmentation 0.555, 17.5% protection, and median risk 48.8. The critical intervention zone contained 65 territories, mean forest integrity 0.431, pressure 0.741, persecution risk 0.600, only 6.2% protection, and median risk 65.2.",
"Under the 300-unit illustrative budget, 86 territories were funded at a total cost of 296 units. Community patrol and anti-persecution action accounted for 61 territories, 183 cost units, and 1,245.63 units of expected risk reduction. Corridor restoration covered nine territories, stewardship or protected-area agreements covered eight, and prey and habitat recovery covered eight. The portfolio concentrated resources on territories where pressure, proximity to roads, and lack of formal protection generated high expected benefit per cost."
]
for x in results6: add_markup_paragraph(doc,x)

full_width_figure(doc,FIGDIR/'figure_6_prioritisation.png',
"Figure 6. Ecological archetypes and auditable prioritisation. (A) Principal-component representation of four synthetic conservation archetypes derived from forest condition, fragmentation, pressure, prey, persecution, and forest loss. (B) Cumulative expected risk reduction as interventions are added in benefit-per-cost order, with the illustrative 300-unit budget indicated. Costs and effects are arbitrary scenario values.")

# ---------- Discussion ----------
add_heading(doc,'Discussion',1)
add_heading(doc,'A nested forest–individual–population–genome system',2)
disc1=[
"The main contribution of this study is not a new estimate of Philippine eagle abundance. It is a transparent architecture for connecting evidence that is usually analysed in separate disciplinary compartments. Forest structure influences prey and movement. Movement determines encounter rates with roads, farms, settlements, traps, and hunters. These exposures influence adult survival and breeding success. Demographic contraction changes effective population size and relatedness. Genomic erosion can then reduce resilience and complicate future recovery. The direction of causality is not exclusively top-down. A declining predator population can alter trophic interactions, and loss of genetically or behaviourally distinct individuals can reduce the range of responses available to the population. Conservation is therefore a coupled system rather than a list of independent threats.",
"The synthetic framework captures this coupling in a form that can be inspected. Every coefficient, threshold, and distribution is visible. This is important because model transparency is part of conservation governance. A high-performing opaque model may still be unsuitable if field teams cannot understand why a territory is prioritised, if local communities cannot challenge incorrect assumptions, or if decision makers cannot distinguish measured data from simulated values."
]
for x in disc1: add_markup_paragraph(doc,x)

add_heading(doc,'Landscape protection must extend beyond nest sites',2)
disc2=[
"The movement benchmark is one of the clearest management signals. A 13 km² core range is only a subset of a 68 km² median home range, and approximately 79% of space-time occurs outside the core (Sutton et al., 2024). Nest protection remains essential, especially during incubation and chick development, but it cannot by itself secure prey, movement corridors, alternative perches, juvenile dispersal routes, or safe access between forest blocks. A protection strategy that ends at the nest buffer may preserve the reproductive platform while allowing the functional territory to degrade.",
"Fragmentation can produce an energetic and exposure penalty. When prey density falls or canopy routes are interrupted, an eagle may have to search farther. Greater flight distance raises energetic expenditure and may lengthen time away from the nest. It can also increase the frequency with which an adult crosses open land, roads, power infrastructure, or areas where persecution occurs. At the physiological level, chronic energetic imbalance can influence body condition, endocrine stress signalling, immune allocation, and reproductive readiness. The present model did not simulate corticosterone, oxidative stress, or metabolomic state, but these are plausible mechanistic links between landscape degradation and reproductive performance that could be tested empirically.",
"The 2035 scenarios illustrate compounding. The annual loss-rate differences were small, but the cumulative difference between integrated protection and high pressure exceeded 880 km². In a species with large territories, hundreds of square kilometres can represent multiple breeding landscapes or critical connections among them. The result supports simultaneous protection of intact forest, restoration of corridors, and land-use agreements in the matrix. It also supports monitoring actual forest structure rather than relying solely on protected-area boundaries."
]
for x in disc2: add_markup_paragraph(doc,x)

add_heading(doc,'Adult survival is a demographic leverage point',2)
disc3=[
"The population scenarios reinforce a general principle of long-lived raptor demography: adult survival has disproportionate influence on population growth. Under the assumptions used here, anti-persecution produced a stronger 40-year population outcome than habitat protection alone, while integrated conservation produced the largest increase. This does not mean that anti-persecution can substitute for habitat. Adult survival protects the reproductive engine that exists today, whereas habitat quality sustains prey, nesting, recruitment, and future carrying capacity. The two interventions act on different but interacting parts of the life cycle.",
"A breeding adult represents accumulated biological capital. It has survived juvenile dispersal, acquired a territory, established a pair bond, learned local prey distributions, and contributed to repeated nesting attempts. Removing such an adult can disrupt the pair, leave dependent young without adequate provisioning, and create a territory vacancy that may not be rapidly filled. In molecular terms, the death also removes one unique diploid genome and one set of rare alleles from a small population. Preventing shooting, trapping, electrocution, collision, and delayed treatment of injuries therefore has demographic and genetic value.",
"The current-pressure median decline to 106 adult females should not be interpreted as a forecast. The assumed rates were not estimated from a national integrated population model. Nevertheless, the contrast among scenarios is useful. It shows that modest changes in adult and pre-breeder survival can dominate long-horizon outcomes, and that uncertainty should motivate sensitivity analysis and better field estimation rather than paralysis."
]
for x in disc3: add_markup_paragraph(doc,x)

add_heading(doc,'Molecular conservation, heterozygosity, and lineage retention',2)
disc4=[
"The genomic component connects conservation action to molecular mechanisms. In a small population, genetic drift changes allele frequencies more strongly because each breeding event represents a larger fraction of the gene pool. Related individuals share longer chromosomal segments inherited from common ancestors. When such individuals reproduce, offspring are more likely to receive identical copies of a segment from both parents, generating runs of homozygosity. Homozygosity can expose recessive variants that are phenotypically silent in heterozygotes. Across many species, inbreeding depression can affect embryo survival, juvenile viability, fertility, morphology, immune function, and resistance to environmental stress, thereby accelerating extinction risk (O'Grady et al., 2006; Spielman et al., 2004).",
"Low mean heterozygosity does not automatically prove current inbreeding depression, and a single genome-wide statistic cannot identify which loci influence fitness. Functional consequences depend on the distribution of deleterious variants, their dominance, local adaptation, gene-by-environment interactions, and the historical effectiveness of purifying selection. The recent Philippine eagle genomic studies provide an important baseline, but conservation interpretation should expand to runs of homozygosity, genomic relatedness, mutational load, sex-linked diversity, immune loci, and temporal sampling (Bacus et al., 2025; Perdon et al., 2026).",
"Mitochondrial and nuclear diversity answer different questions. Seventeen mitochondrial haplotypes suggest that multiple maternal lineages remain, even though nucleotide diversity is low. A breeding programme focused only on nuclear mean kinship could inadvertently lose a rare maternal haplotype. Conversely, retaining every haplotype without controlling nuclear relatedness could increase inbreeding. A multi-objective strategy should preserve founder and maternal-line representation while minimising mean kinship and avoiding overproduction by genetically overrepresented individuals.",
"The 23.9% optimisation gain demonstrates the mathematical value of formal assignment, but the objective function was simplified. Real pairing must also incorporate age, health, reproductive competence, behavioural compatibility, existing pair bonds, welfare, facility capacity, disease risk, and legal or ethical constraints. Geographic origin should not be used as an automatic proxy for genetic distance. The risk of outbreeding depression, local adaptation, and disruption of population structure should be assessed from genomic and ecological evidence before translocation or cross-population breeding. Conservation genomics is most useful when it improves decisions, not when it merely produces sequence data (Shafer et al., 2015; Supple & Shapiro, 2018)."
]
for x in disc4: add_markup_paragraph(doc,x)

add_heading(doc,'Interpretable machine learning and decision support',2)
disc5=[
"The fledging model was intentionally interpretable and achieved moderate discrimination. In conservation, moderate predictive performance can still be useful when it identifies data gaps, ranks inspection priorities, or supports adaptive monitoring. However, model evaluation must respect spatial and temporal structure. Randomly dividing observations from the same territories can overstate generalisability because repeated records share latent habitat and pair characteristics. A field implementation should use blocked cross-validation by territory, landscape, and time, and should model detection probability, missing nests, observer effort, and repeated measures.",
"Human pressure emerged as the largest negative coefficient because the synthetic outcome function encoded it strongly. This illustrates a key distinction between simulation validation and empirical discovery. The model correctly recovered the designed data-generating process, but the coefficient magnitude cannot be cited as a field effect. A valuable use of synthetic data is precisely to make this distinction visible. It allows researchers to test whether a pipeline can recover known signals, quantify calibration, and detect leakage before sensitive or costly field data are introduced.",
"The clustering and portfolio modules demonstrate how analytical outputs can be converted into decisions. The archetypes are more actionable than a single risk continuum because they suggest different intervention logics. A fragmented pressure front may need corridors and negotiated land use, whereas a critical intervention zone may first require immediate community patrols, conflict reduction, and road-interface management. The portfolio model made costs and assumed effects explicit, so stakeholders can replace them with locally derived values and observe how priorities change."
]
for x in disc5: add_markup_paragraph(doc,x)

add_heading(doc,'Management implications',2)
disc6=[
"The integrated results support five mutually reinforcing actions. First, prevent avoidable adult mortality through community-based anti-persecution programmes, rapid veterinary response, safe power infrastructure, and targeted work near roads and forest edges. Second, protect functional territories rather than only nest trees, including foraging habitat, movement corridors, and buffer landscapes. Third, restore prey-supporting forest structure and canopy connectivity where fragmentation increases ranging demand. Fourth, establish a secure national conservation-genomics programme that links wild and ex situ individuals through validated identity, pedigree, relatedness, and lineage data. Fifth, use transparent scenario models as adaptive tools, updating them whenever new survival, reproduction, forest, or genomic observations become available.",
"Community participation is not an optional communication layer added after modelling. It is a causal component of conservation success. Local residents, Indigenous communities, forest guards, landowners, and municipal authorities influence hunting risk, reporting, land conversion, rescue response, and acceptance of protected or stewardship agreements. Ecological models should therefore incorporate social variables, intervention feasibility, and local knowledge rather than treating human pressure as an undifferentiated index.",
"Data security is also part of species protection. Fine-scale nest and telemetry data should be stored under role-based access, with public products aggregated to resolutions that prevent re-identification. Genomic samples require chain-of-custody, permits, benefit sharing, and governance consistent with Philippine law and institutional ethics. Reproducibility should mean that methods and non-sensitive summaries are open, not that sensitive locations are exposed."
]
for x in disc6: add_markup_paragraph(doc,x)

add_heading(doc,'Limitations and future empirical programme',2)
disc7=[
"Every territory, breeding outcome, demographic transition, cost, intervention effect, and genomic profile in this study was simulated. The island allocation was not a census. The reproductive model omitted nest detection, observer variation, pair turnover, individual age, disease, contaminant exposure, typhoon damage, prey-specific dynamics, and density dependence. The PVA used illustrative survival and reproduction parameters and did not include catastrophes, spatial dispersal, carrying capacity, sex-specific pairing limitation, or inbreeding depression. The genomic pairing score did not model real relatedness or functional variants. The portfolio assumed linear and immediate intervention benefits.",
"A national empirical programme should integrate verified nest occupancy and productivity, marked-bird survival, rescue and mortality records, telemetry with strict security, forest structure, roads and power infrastructure, prey abundance, climate and typhoon exposure, community-reported persecution, and genomic relatedness. Integrated population models could combine count, occupancy, productivity, and survival data. Before–after–control–impact designs could estimate intervention effects. Genomic sampling should be longitudinal and representative across islands, while avoiding overinterpretation of small sample sizes.",
"At the molecular level, future work could relate individual heterozygosity and runs of homozygosity to semen quality, hatchability, embryonic loss, immune phenotype, pathogen burden, endocrine stress, and survival, with rigorous control for age, environment, and pedigree. Such work would test whether genomic diversity is already affecting fitness and identify which management actions have the greatest biological return."
]
for x in disc7: add_markup_paragraph(doc,x)

# ---------- Conclusion ----------
add_heading(doc,'Conclusion',1)
conclusion=("A literature-benchmarked synthetic framework can unite the principal scales of Philippine eagle conservation, from forest integrity and movement exposure to breeding, adult survival, population trajectory, and genomic diversity. The model reproduced published central benchmarks while keeping all sensitive individual-level information artificial. Its strongest systems-level message is that adult survival, functional landscape protection, habitat connectivity, prey support, and lineage-aware genetic management are complementary. Nest protection without safe ranging habitat is incomplete, habitat protection without control of persecution may fail to retain breeding adults, and demographic recovery without genetic stewardship may erode long-term resilience. The notebook provides a transparent platform for hypothesis testing, training, and stakeholder discussion. Its simulated outputs are not biological discoveries or an official recovery plan, but they define an auditable structure into which secure empirical data can be progressively incorporated.")
add_markup_paragraph(doc,conclusion)

add_heading(doc,'Data and Code Availability',1)
add_markup_paragraph(doc,"The accompanying Jupyter notebook generates all synthetic records, validation summaries, scenario projections, and figures. The executed workflow exports 27 files, including 10 CSV datasets and 16 high-resolution figures. All records are synthetic and contain no real nest locations, telemetry coordinates, identities, pedigrees, or sequences. The random seed is 20260716.")

add_heading(doc,'Author Contributions',1)
add_markup_paragraph(doc,"M.I.R.P. conceived the study, curated the literature benchmarks, designed and implemented the synthetic simulations, performed the analyses, generated the figures, interpreted the results, and wrote the manuscript.")

add_heading(doc,'Funding',1)
add_markup_paragraph(doc,"No external funding was received for this synthetic methods study.")

add_heading(doc,'Competing Interests',1)
add_markup_paragraph(doc,"The author declares no competing interests.")

add_heading(doc,'Ethics and Conservation-Security Statement',1)
add_markup_paragraph(doc,"No animals were captured, handled, sampled, tracked, or experimentally manipulated. No human participants were enrolled. No real nest coordinates, telemetry records, individual identities, pedigrees, or genotypes were used. The manuscript should not be used to infer or publish the location of any Philippine eagle.")

# ---------- References ----------
add_heading(doc,'References',1)
refs=[
("Bacus, M. G., et al. (2025). Mitogenomic diversity and population structure of the Philippine eagle. Ecology and Evolution.","10.1002/ece3.72572","PMID: 41356510."),
("Brook, B. W., O'Grady, J. J., Chapman, A. P., Burgman, M. A., Akçakaya, H. R., & Frankham, R. (2000). Predictive accuracy of population viability analysis in conservation biology. Nature, 404, 385–387.","10.1038/35002056",None),
("Bueser, G. L. L., Bueser, K. G., Afan, D. S., Salvador, D. I., Grier, J. W., Kennedy, R. S., & Miranda, H. C. (2003). Distribution and nesting density of the Philippine eagle Pithecophaga jefferyi on Mindanao Island, Philippines: What do we know after 100 years? Ibis, 145, 130–135.","10.1046/j.1474-919X.2003.00131.x",None),
("Christin, S., Hervet, É., & Lecomte, N. (2019). Applications for deep learning in ecology. Methods in Ecology and Evolution, 10, 1632–1644.","10.1111/2041-210X.13256",None),
("Daipan, B. P. O. (2021). Forest loss in terrestrial Key Biodiversity Areas in the Philippines, 2001–2019. Journal of Threatened Taxa, 13, 20019–20032.","10.11609/jott.6904.13.13.20019-20032",None),
("Frankham, R., Bradshaw, C. J. A., & Brook, B. W. (2014). Genetics in conservation management: Revised recommendations for the 50/500 rules, Red List criteria and population viability analyses. Biological Conservation, 170, 56–63.","10.1016/j.biocon.2013.12.036",None),
("Gibson, L., Lee, T. M., Koh, L. P., Brook, B. W., Gardner, T. A., Barlow, J., Peres, C. A., Bradshaw, C. J. A., Laurance, W. F., Lovejoy, T. E., & Sodhi, N. S. (2011). Primary forests are irreplaceable for sustaining tropical biodiversity. Nature, 478, 378–381.","10.1038/nature10425",None),
("Haddad, N. M., Brudvig, L. A., Clobert, J., Davies, K. F., Gonzalez, A., Holt, R. D., Lovejoy, T. E., Sexton, J. O., Austin, M. P., Collins, C. D., Cook, W. M., Damschen, E. I., Ewers, R. M., Foster, B. L., Jenkins, C. N., King, A. J., Laurance, W. F., Levey, D. J., Margules, C. R., et al. (2015). Habitat fragmentation and its lasting impact on Earth's ecosystems. Science Advances, 1, e1500052.","10.1126/sciadv.1500052",None),
("Ibañez, J. C., Miranda, H. C., Balaquit-Ibañez, G., Afan, D. S., & Kennedy, R. S. (2003). Notes on the breeding behavior of a Philippine eagle pair at Mount Sinaka, Central Mindanao. Wilson Bulletin, 115, 333–336.","10.1676/01-054",None),
("Kardos, M., Armstrong, E. E., Fitzpatrick, S. W., Hauser, S., Hedrick, P. W., Miller, J. M., Tallmon, D. A., & Funk, W. C. (2021). The crucial role of genome-wide genetic variation in conservation. Proceedings of the National Academy of Sciences of the United States of America, 118, e2104642118.","10.1073/pnas.2104642118",None),
("Newbold, T., Hudson, L. N., Hill, S. L. L., Contu, S., Lysenko, I., Senior, R. A., Börger, L., Bennett, D. J., Choimes, A., Collen, B., Day, J., De Palma, A., Díaz, S., Echeverria-Londoño, S., Edgar, M. J., Feldman, A., Garon, M., Harrison, M. L. K., Alhusseini, T., et al. (2015). Global effects of land use on local terrestrial biodiversity. Nature, 520, 45–50.","10.1038/nature14324",None),
("O'Grady, J. J., Brook, B. W., Reed, D. H., Ballou, J. D., Tonkyn, D. W., & Frankham, R. (2006). Realistic levels of inbreeding depression strongly affect extinction risk in wild populations. Biological Conservation, 133, 42–51.","10.1016/j.biocon.2006.05.016",None),
("Ong, P. S., et al. (2011). DNA barcodes of Philippine accipitrids. Molecular Ecology Resources.","10.1111/j.1755-0998.2010.02928.x","PMID: 21429130."),
("Perdon, J. P. M., et al. (2026). Genome-wide diversity and demographic history of the critically endangered Philippine eagle. BMC Genomics.","10.1186/s12864-026-12859-9","PMID: 42010461."),
("Perez, G. J., et al. (2020). Forest-cover change in the Sierra Madre Biodiversity Corridor, Philippines. Forests, 11, 1071.","10.3390/f11101071","FAO AGRIS record."),
("Roberts, D. R., Bahn, V., Ciuti, S., Boyce, M. S., Elith, J., Guillera-Arroita, G., Hauenstein, S., Lahoz-Monfort, J. J., Schröder, B., Thuiller, W., Warton, D. I., Wintle, B. A., Hartig, F., & Dormann, C. F. (2017). Cross-validation strategies for data with temporal, spatial, hierarchical, or phylogenetic structure. Ecography, 40, 913–929.","10.1111/ecog.02881",None),
("Shafer, A. B. A., Wolf, J. B. W., Alves, P. C., Bergström, L., Bruford, M. W., Brännström, I., Colling, G., Dalén, L., De Meester, L., Ekblom, R., Fawcett, K. D., Fior, S., Hajibabaei, M., Hill, J. A., Hoezel, A. R., Höglund, J., Jensen, E. L., Krause, J., Kristensen, T. N., et al. (2015). Genomics and the challenging translation into conservation practice. Trends in Ecology & Evolution, 30, 78–87.","10.1016/j.tree.2014.11.009",None),
("Spielman, D., Brook, B. W., & Frankham, R. (2004). Most species are not driven to extinction before genetic factors impact them. Proceedings of the National Academy of Sciences of the United States of America, 101, 15261–15264.","10.1073/pnas.0403809101",None),
("Steiner, C. C., Putnam, A. S., Hoeck, P. E. A., & Ryder, O. A. (2013). Conservation genomics of threatened animal species. Annual Review of Animal Biosciences, 1, 261–281.","10.1146/annurev-animal-031412-103636",None),
("Supple, M. A., & Shapiro, B. (2018). Conservation of biodiversity in the genomics era. Genome Biology, 19, 131.","10.1186/s13059-018-1520-3",None),
("Sutton, L. J., et al. (2023). Priority conservation areas and a global population estimate for the critically endangered Philippine eagle. Animal Conservation.","10.1111/acv.12854",None),
("Sutton, L. J., et al. (2024). Space-time home-range estimates and resource selection for the critically endangered Philippine eagle on Mindanao. Ibis.","10.1111/ibi.13233",None),
("Tuia, D., Kellenberger, B., Beery, S., Costelloe, B. R., Zuffi, S., Risse, B., Mathis, A., Mathis, M. W., van Langevelde, F., Burghardt, T., Kays, R., Klinck, H., Wikelski, M., Couzin, I. D., van Horn, G., Crofoot, M. C., Stewart, C. V., & Berger-Wolf, T. (2022). Perspectives in machine learning for wildlife conservation. Nature Communications, 13, 792.","10.1038/s41467-022-27980-y",None),
("Watson, J. E. M., Evans, T., Venter, O., Williams, B., Tulloch, A., Stewart, C., Thompson, I., Ray, J. C., Murray, K., Salazar, A., McAlpine, C., Potapov, P., Walston, J., Robinson, J. G., Painter, M., Wilkie, D., Filardi, C., Laurance, W. F., Houghton, R. A., et al. (2018). The exceptional value of intact forest ecosystems. Nature Ecology & Evolution, 2, 599–610.","10.1038/s41559-018-0490-x",None),
]
for t,d,n in refs: add_reference(doc,t,d,n)

add_markup_paragraph(doc,"Database provenance statement. The Philippine eagle benchmark articles were selected from PubMed records, FAO AGRIS records, or journals represented in USDA AGRICOLA resources, as documented in Supplementary Table S1. General ecological and genomic references were restricted to peer-reviewed journals represented in PubMed, AGRIS, or AGRICOLA. Indexing should be rechecked at formal submission because database coverage can change.",style='References',first_line=False,space_after=4)

add_heading(doc,'Author Preprints Included for Methodological Context',1)
add_markup_paragraph(doc,"The following preprints are included because the author requested representation of his bioRxiv, ChemRxiv, and medRxiv work. They document continuity in transparent simulation, systems modelling, and reproducible computational design. They are not presented as peer-reviewed or as direct evidence for Philippine eagle biology.",first_line=False)
preprints=[
("Petalcorin, M. I. R. (2026a). Computational synthetic inner membrane reveals cardiolipin–leak control of ATP output [Preprint]. bioRxiv.","10.64898/2026.02.25.708092",None),
("Petalcorin, M. I. R. (2026b). Reproducible genomic medicine stack for population structure, statistical genetics, tumour signatures, clonality, and AI variant prioritisation [Preprint]. ChemRxiv.","10.26434/chemrxiv.15000285/v1",None),
("Petalcorin, M. I. R. (2026c). An end-to-end synthetic oncology clinical trial framework integrating radiographic response, circulating tumour DNA, safety, and survival for decision-oriented clinical data science [Preprint]. medRxiv.","10.64898/2026.04.07.26350297",None),
]
for t,d,n in preprints: add_reference(doc,t,d,n)

# ---------- Supplementary single-column section ----------
supp=doc.add_section(WD_SECTION_START.NEW_PAGE); set_columns(supp,1)
add_heading(doc,'Supplementary Material',1)
add_markup_paragraph(doc,"All tables are placed in this supplementary section to preserve the figure-centred, two-column flow of the main manuscript. All individual-level records are synthetic.",first_line=False)

bench=pd.read_csv(OUTDIR/'literature_benchmarks.csv')
bench=bench[['domain','study','metric','value','unit','uncertainty','doi','PubMed','AGRIS','AGRICOLA_journal']]
bench.columns=['Domain','Study','Metric','Value','Unit','Uncertainty','DOI','PubMed','AGRIS','AGRICOLA journal']
add_table_from_df(doc,bench,'Supplementary Table S1. Literature benchmarks, numerical targets, and index provenance.',font_size=6.2,formats={'Value':lambda x: f'{x:g}'})

val=pd.read_csv(OUTDIR/'benchmark_validation.csv')
val.columns=['Metric','Published target','Synthetic value','Unit','Relative error, %']
add_table_from_df(doc,val,'Supplementary Table S2. Calibration of synthetic central values to published targets.',font_size=7.1,formats={'Published target':lambda x:f'{x:.3f}','Synthetic value':lambda x:f'{x:.3f}','Relative error, %':lambda x:f'{x:.3f}'})

breed_summary=pd.DataFrame({
    'Metric':['Total territory-year records','Active attempts','Eggs per active attempt','Mean incubation, days','Mean female incubation share, %','Hatching success, %','Fledging success per attempt, %','Test ROC-AUC','Brier score','Fivefold CV ROC-AUC'],
    'Value':[3528,2950,1.0,58.006,74.061,62.373,37.932,0.674,0.225,0.656]
})
add_table_from_df(doc,breed_summary,'Supplementary Table S3. Synthetic breeding and model-performance summary.',font_size=7.3,formats={'Value':lambda x:f'{x:.3f}'})

pva=pd.read_csv(OUTDIR/'pva_scenario_summary.csv')
pva.columns=['Scenario','Median females, year 40','Lower 95%','Upper 95%','Probability below 50, %','Probability zero, %']
add_table_from_df(doc,pva,'Supplementary Table S4. Forty-year stochastic population-viability outcomes.',font_size=7.0,formats={c:(lambda x:f'{x:.2f}') for c in pva.columns[1:]})

forest=pd.read_csv(OUTDIR/'forest_scenario_projection.csv')
forest=forest[forest['year']==2035][['scenario','remaining_aoh_km2','loss_from_2026_km2']]
forest.columns=['Scenario','Remaining AOH in 2035, km²','Loss from 2026, km²']
add_table_from_df(doc,forest,'Supplementary Table S5. Area-of-habitat scenario outcomes in 2035.',font_size=7.2,formats={forest.columns[1]:lambda x:f'{x:,.1f}',forest.columns[2]:lambda x:f'{x:,.1f}'})

gen_summary=pd.DataFrame({
    'Metric':['Nuclear individuals','Mean heterozygosity','Minimum heterozygosity','Mitochondrial individuals','Mitochondrial haplotypes','Mean nucleotide-diversity proxy','Haplotype diversity','Optimised mean pair score','Random mean pair score','Optimisation gain, %'],
    'Value':[35,0.000309,0.000207,32,17,0.000540,0.947581,1.983,1.600,23.9]
})
add_table_from_df(doc,gen_summary,'Supplementary Table S6. Synthetic conservation-genomic summary.',font_size=7.2,formats={'Value':lambda x:(f'{x:.6f}' if abs(x)<0.01 else f'{x:.3f}')})

profiles=pd.DataFrame([
['Relatively secure forest',119,0.773,0.201,0.279,0.716,0.234,57.143,19.800],
['Watch-list landscape',105,0.706,0.290,0.521,0.576,0.404,33.333,38.474],
['Fragmented pressure front',103,0.472,0.555,0.398,0.464,0.377,17.476,48.779],
['Critical intervention zone',65,0.431,0.597,0.741,0.370,0.600,6.154,65.166],
],columns=['Archetype','n','Forest','Fragmentation','Pressure','Prey','Persecution','Protected, %','Median risk'])
add_table_from_df(doc,profiles,'Supplementary Table S7. Ecological profiles of the four synthetic territory archetypes.',font_size=6.8,formats={c:(lambda x:f'{x:.3f}') for c in profiles.columns[2:]})

portfolio_summary=pd.DataFrame([
['Community patrol and anti-persecution',61,183.0,1245.63],
['Corridor restoration',9,45.0,184.57],
['Stewardship or protected-area agreement',8,36.0,137.19],
['Prey and habitat recovery',8,32.0,125.53],
],columns=['Recommended action','Territories','Cost units','Expected risk reduction'])
add_table_from_df(doc,portfolio_summary,'Supplementary Table S8. Illustrative intervention portfolio under a 300-unit budget.',font_size=7.0,formats={'Cost units':lambda x:f'{x:.1f}','Expected risk reduction':lambda x:f'{x:.2f}'})

# finalize all sections margins/header/footer
for s in doc.sections:
    s.page_width=Cm(21.0); s.page_height=Cm(29.7)
    s.top_margin=Inches(0.55); s.bottom_margin=Inches(0.55); s.left_margin=Inches(0.62); s.right_margin=Inches(0.62)
    # Unlink section headers and footers before writing, otherwise python-docx
    # appends repeated PAGE fields to a shared linked footer.
    s.header.is_linked_to_previous = False
    s.footer.is_linked_to_previous = False
    hp=s.header.paragraphs[0]
    hp.clear()
    hp.text='PETALCORIN  |  PHILIPPINE EAGLE CONSERVATION ANALYTICS'
    hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs: r.font.name='Arial'; r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string('666666')
    fp=s.footer.paragraphs[0]
    fp.clear()
    add_page_number(fp)

# metadata
props=doc.core_properties
props.title='Integrated synthetic conservation analytics for the Philippine eagle'
props.subject='Philippine eagle conservation, habitat, demography, and genomics'
props.author='Mark Ihrwell R. Petalcorin'
props.keywords='Philippine eagle; conservation genomics; population viability; habitat fragmentation; synthetic data'
props.comments='All individual-level records are synthetic.'

doc.save(DOCX)
print(DOCX)
